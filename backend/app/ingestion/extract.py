"""Table-aware content extraction for PDF and DOCX.

The old prototype only read paragraph text and silently dropped every table.
Real BRDs/SDDs put critical data (field mappings, attribute lists, campaign
matrices) in tables, so every block here is tagged with its kind
("heading" | "paragraph" | "table") and tables are rendered as markdown so
the LLM downstream can read them as structured data, not noise.
"""

import pathlib

import fitz
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def iter_block_items(doc: Document):
    """Yield paragraphs and tables in the order they appear in the document body."""
    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def extract_blocks_from_docx(file_path: str) -> list[dict]:
    doc = Document(file_path)
    blocks = []
    table_index = 0
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            style = item.style.name.lower() if item.style else ""
            kind = "heading" if ("heading" in style or style.startswith("title")) else "paragraph"
            blocks.append({"kind": kind, "text": text, "style": style})
        elif isinstance(item, Table):
            table_index += 1
            md = table_to_markdown(item)
            if md:
                blocks.append({
                    "kind": "table",
                    "text": f"[TABLE {table_index}]\n{md}",
                    "style": "table",
                })
    return blocks


def extract_blocks_from_pdf(file_path: str) -> list[dict]:
    doc = fitz.open(file_path)
    blocks: list[dict] = []
    body_sizes = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block["type"] == 0:
                for line in block["lines"]:
                    for span in line["spans"]:
                        body_sizes.append(span["size"])
    heading_threshold = (sum(body_sizes) / len(body_sizes) * 1.2) if body_sizes else 999

    table_index = 0
    for page in doc:
        try:
            found_tables = page.find_tables()
        except Exception:
            found_tables = None
        table_bboxes = []
        if found_tables and found_tables.tables:
            for t in found_tables.tables:
                table_index += 1
                try:
                    rows = t.extract()
                except Exception:
                    rows = []
                if rows:
                    header = [str(c or "") for c in rows[0]]
                    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
                    for row in rows[1:]:
                        lines.append("| " + " | ".join(str(c or "") for c in row) + " |")
                    blocks.append({
                        "kind": "table",
                        "text": f"[TABLE {table_index}]\n" + "\n".join(lines),
                        "style": "table",
                    })
                table_bboxes.append(t.bbox)

        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            bbox = block.get("bbox")
            if bbox and any(_bbox_inside(bbox, tb) for tb in table_bboxes):
                continue
            for line in block["lines"]:
                line_text = ""
                max_size = 0.0
                for span in line["spans"]:
                    line_text += span["text"]
                    max_size = max(max_size, span["size"])
                line_text = line_text.strip()
                if not line_text:
                    continue
                kind = "heading" if (max_size >= heading_threshold and len(line_text) < 100) else "paragraph"
                blocks.append({"kind": kind, "text": line_text, "style": f"size:{max_size:.1f}"})
    doc.close()
    return blocks


def _bbox_inside(bbox, table_bbox, tolerance: float = 2.0) -> bool:
    x0, y0, x1, y1 = bbox
    tx0, ty0, tx1, ty1 = table_bbox
    return x0 >= tx0 - tolerance and y0 >= ty0 - tolerance and x1 <= tx1 + tolerance and y1 <= ty1 + tolerance


def extract_blocks(file_path: str) -> dict:
    path = pathlib.Path(file_path)
    extension = path.suffix.lower()
    if extension == ".docx":
        blocks = extract_blocks_from_docx(file_path)
    elif extension == ".pdf":
        blocks = extract_blocks_from_pdf(file_path)
    else:
        return {"error": f"Unsupported file type '{extension}'. Please upload PDF or DOCX."}

    table_count = sum(1 for b in blocks if b["kind"] == "table")
    heading_count = sum(1 for b in blocks if b["kind"] == "heading")
    return {
        "blocks": blocks,
        "block_count": len(blocks),
        "table_count": table_count,
        "heading_count": heading_count,
        "file_type": extension.lstrip("."),
    }
